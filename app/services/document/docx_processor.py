from app.processors.base_processor import BaseDocumentProcessor
from typing import List
from docx import Document
import re
import logging

class DocxProcessor(BaseDocumentProcessor):
    """DOCX 파일 처리기"""
    
    def process(self, file_path: str) -> List[str]:
        """
        DOCX 파일에서 향상된 텍스트 추출
        
        Args:
            file_path (str): DOCX 파일 경로
            
        Returns:
            List[str]: 추출된 텍스트 청크 목록
        """
        try:
            doc = Document(file_path)
            paragraphs_text = []
            
            # 문서 메타데이터 수집
            metadata = self._extract_document_metadata(doc)
            if metadata:
                paragraphs_text.append(f"[문서 정보]\n{metadata}")
            
            # 구조화된 텍스트 추출
            current_section = ""
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                    
                # 스타일 정보 확인
                style_name = para.style.name if para.style else ""
                
                # 헤더 스타일 감지
                if any(header_style in style_name.lower() for header_style in ['heading', '제목']):
                    current_section = text
                    paragraphs_text.append(f"\n[섹션: {text}]")
                else:
                    # 일반 텍스트에 섹션 컨텍스트 추가
                    if current_section:
                        text = f"{text}"
                    paragraphs_text.append(text)
            
            # 표 추출 (구조 정보 포함)
            for i, table in enumerate(doc.tables):
                table_data = self._extract_table_with_structure(table, i)
                if table_data:
                    paragraphs_text.append(table_data)
            
            # 헤더/푸터 처리
            header_footer_text = self._extract_headers_footers(doc)
            if header_footer_text:
                paragraphs_text.insert(1, header_footer_text)  # 메타데이터 다음에 삽입
            
            # 전체 텍스트 구성
            full_text = "\n\n".join(paragraphs_text)
            
            # 향상된 청킹 (문서 구조 고려)
            chunks = self._smart_chunk_text(full_text)
            
            logging.info(f"DOCX 처리 완료: {file_path}, {len(chunks)}개 청크 생성")
            return chunks
            
        except Exception as e:
            logging.error(f"DOCX 처리 중 오류 발생: {e}")
            return [f"DOCX 파일 처리 중 오류가 발생했습니다: {str(e)}"]
    
    def _extract_document_metadata(self, doc) -> str:
        """문서 메타데이터 추출"""
        metadata_parts = []
        
        if doc.core_properties.title:
            metadata_parts.append(f"제목: {doc.core_properties.title}")
        
        if doc.core_properties.author:
            metadata_parts.append(f"작성자: {doc.core_properties.author}")
            
        if doc.core_properties.created:
            metadata_parts.append(f"생성일: {doc.core_properties.created.strftime('%Y-%m-%d')}")
            
        if doc.core_properties.subject:
            metadata_parts.append(f"주제: {doc.core_properties.subject}")
            
        if doc.core_properties.keywords:
            metadata_parts.append(f"키워드: {doc.core_properties.keywords}")
        
        return "\n".join(metadata_parts) if metadata_parts else ""
    
    def _extract_table_with_structure(self, table, table_index: int) -> str:
        """표 구조 정보와 함께 추출"""
        try:
            table_text = f"\n[표 {table_index + 1}]\n"
            
            # 헤더 행 식별 시도
            if table.rows:
                first_row = table.rows[0]
                header_cells = [cell.text.strip() for cell in first_row.cells]
                
                if any(header_cells):  # 헤더가 있는 경우
                    table_text += "헤더: " + " | ".join(header_cells) + "\n"
                    
                    # 데이터 행들 처리
                    for i, row in enumerate(table.rows[1:], 1):
                        row_data = [cell.text.strip() for cell in row.cells]
                        if any(row_data):
                            table_text += f"행{i}: " + " | ".join(row_data) + "\n"
                else:
                    # 헤더가 없는 경우
                    for i, row in enumerate(table.rows):
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            table_text += f"행{i+1}: " + " | ".join(row_data) + "\n"
            
            return table_text if len(table_text) > f"\n[표 {table_index + 1}]\n" else ""
            
        except Exception as e:
            logging.warning(f"표 {table_index} 추출 중 오류: {e}")
            return ""
    
    def _extract_headers_footers(self, doc) -> str:
        """헤더/푸터 추출"""
        header_footer_parts = []
        
        try:
            # 섹션별 헤더/푸터 처리
            for section in doc.sections:
                # 헤더 추출
                if section.header:
                    header_text = ""
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            header_text += para.text.strip() + " "
                    
                    if header_text.strip():
                        header_footer_parts.append(f"헤더: {header_text.strip()}")
                
                # 푸터 추출
                if section.footer:
                    footer_text = ""
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            footer_text += para.text.strip() + " "
                    
                    if footer_text.strip():
                        header_footer_parts.append(f"푸터: {footer_text.strip()}")
        
        except Exception as e:
            logging.warning(f"헤더/푸터 추출 중 오류: {e}")
        
        return "\n".join(header_footer_parts) if header_footer_parts else ""
    
    def _smart_chunk_text(self, text: str, max_chunk_size: int = 800, overlap: int = 150) -> List[str]:
        """
        문서 구조를 고려한 스마트 청킹
        """
        if not text or len(text) <= max_chunk_size:
            return [text] if text else []
        
        chunks = []
        
        # 섹션별로 먼저 분할
        sections = re.split(r'\n\[섹션:[^\]]+\]\n', text)
        
        for section in sections:
            if not section.strip():
                continue
                
            # 섹션이 크기 제한을 초과하는 경우 추가 분할
            if len(section) <= max_chunk_size:
                chunks.append(section.strip())
            else:
                # 기존 청킹 로직 사용 (개선된 파라미터로)
                section_chunks = self.chunk_text(section, max_chunk_size, overlap)
                chunks.extend(section_chunks)
        
        return [chunk for chunk in chunks if chunk.strip()]